from fastapi.responses import FileResponse

from transformers import BartTokenizer, BartForConditionalGeneration
from docx import Document
from fastapi import FastAPI, UploadFile, File
from io import BytesIO


app = FastAPI()

model_name = 'facebook/bart-large-cnn'
tokenizer = BartTokenizer.from_pretrained(model_name)
model = BartForConditionalGeneration.from_pretrained(model_name)

def summarize_document(document):
    paragraphs = document.split('\n')
    long_paragraphs = [p for p in paragraphs if len(p.split()) > 20]
    summaries = []
    
    max_summary_length = 0.5 * len(document.split())

    for paragraph in long_paragraphs:
        inputs = tokenizer.encode(paragraph, return_tensors='pt')

        summary_ids = model.generate(inputs, max_length=int(max_summary_length), num_beams=4, early_stopping=False)
        summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

        summaries.append(summary)

    return summaries

@app.get("/")
def welcome():
    return {"message": "Welcome to text summarization"}


def generate_summary_doc(summaries):
    doc = Document()

    for i, summary in enumerate(summaries, 1):
        doc.add_paragraph(f"Summary {i}: {summary}")

    return doc

@app.post("/summarize")
async def summarize(file: UploadFile = File(...)):
    document_bytes = await file.read()
    document = Document(BytesIO(document_bytes))
    text = [paragraph.text for paragraph in document.paragraphs]
    data = '\n'.join(text)
    summaries = summarize_document(data)

    summary_doc = generate_summary_doc(summaries)
    output_path = "generated_summaries.docx"
    summary_doc.save(output_path)

    return FileResponse(output_path, filename="generated_summaries.docx")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)

