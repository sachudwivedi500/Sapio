from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from docx import Document
import re

app = FastAPI()

class DocumentRequest(BaseModel):
    document_path: str

class DocumentResponse(BaseModel):
    data: str

def generate_summary(paragraphs):
    # Placeholder function for generating summaries
    summaries = []
    for paragraph in paragraphs:
        # Perform your summary generation logic here
        summary = "Summary for paragraph: " + paragraph
        summaries.append(summary)
    return summaries

@app.post("/upload_document")
async def upload_document(file: UploadFile = File(...)):
    # Save the uploaded DOC file
    file_path = f"C:/Users/User/Downloads/Company1 - IT - 2023.docx"
    with open(file_path, "wb") as f:
        f.write(await file.read())

    return {"message": "Document uploaded successfully"}

@app.post("/extract_text")
def extract_text_from_document(request: DocumentRequest) -> DocumentResponse:
    # Open the Word document
    doc = Document(request.document_path)

    # Extract the text from paragraphs in the document
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    # Join the extracted text into a single string
    data = '\n'.join(text)

    return DocumentResponse(data=data)

@app.post("/identify_paragraphs")
def identify_paragraphs(request: DocumentRequest):
    doc = Document(request.document_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    data = '\n'.join(text)

    # Split the document into paragraphs using a regex pattern
    paragraphs = re.split(r'\n\s*\n', data)

    # Remove extra whitespace and special characters from each paragraph
    paragraphs = [re.sub(r'\s+', ' ', paragraph.strip()) for paragraph in paragraphs]

    identified_paragraphs = paragraphs

    return {"identified_paragraphs": identified_paragraphs}

@app.post("/generate_summaries")
def generate_summaries(request: DocumentRequest):
    doc = Document(request.document_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    data = '\n'.join(text)

    paragraphs = re.split(r'\n\s*\n', data)
    paragraphs = [re.sub(r'\s+', ' ', paragraph.strip()) for paragraph in paragraphs]
    paragraphs= str(paragraphs)

    summaries = generate_summary(paragraphs)  # Implement your own summary generation function

    return { summaries}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)